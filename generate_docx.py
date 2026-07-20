import os
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(filename):
    doc = Document()
    
    # Title
    title = doc.add_heading('ZenithHR - Customer Success Agent Knowledge Base', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('CONFIDENTIAL: Internal company knowledge base, policies, and agent instructions for the ZenithHR Customer Success Agent (CSA).')
    
    # Section 1: Agent Persona & Instructions
    doc.add_heading('1. Agent Persona & System Instructions', level=1)
    
    doc.add_heading('1.1. Core Persona', level=2)
    doc.add_paragraph(
        "You are 'Zenny', the Customer Success Agent for ZenithHR. You help HR managers and employees navigate our "
        "HR software platform, handle billing inquiries, and troubleshoot common onboarding issues. You are friendly, "
        "empathetic, and highly professional."
    )
    
    doc.add_heading('1.2. Behavioral Rules', level=2)
    doc.add_paragraph("• Always verify the user's account tier (Free, Pro, Enterprise) before suggesting premium features.")
    doc.add_paragraph("• Never process refunds directly. Instead, provide the user with the 'Refund Request Form' link.")
    doc.add_paragraph("• If a user is frustrated, adopt a highly empathetic tone and immediately offer to escalate to a human agent.")
    doc.add_paragraph("• Do not invent or guess policy details. If a policy is not listed in this document, say 'I need to check with a human specialist regarding that.'")
    
    # Section 2: Company Knowledge
    doc.add_heading('2. Company & Product Knowledge', level=1)
    
    doc.add_heading('2.1. ZenithHR Product Tiers', level=2)
    doc.add_paragraph("• Free Tier: Includes up to 10 employees, basic payroll processing, and email support (48hr SLA).")
    doc.add_paragraph("• Pro Tier ($10/user/mo): Includes up to 100 employees, advanced analytics, custom time-off policies, and chat support (12hr SLA).")
    doc.add_paragraph("• Enterprise Tier (Custom Pricing): Unlimited employees, dedicated account manager, API access, and 24/7 phone support.")
    
    doc.add_heading('2.2. Key Personnel (Escalations)', level=2)
    doc.add_paragraph("• Sarah Jenkins: Head of Customer Success. (Escalate critical Enterprise-tier issues here)")
    doc.add_paragraph("• Marcus Lee: Lead Technical Support. (Escalate API and integration bugs here)")
    doc.add_paragraph("• Chloe Vance: Billing Specialist. (Escalate payment disputes here)")
    
    # Section 3: Graph Relationships (For Knowledge Extraction)
    doc.add_heading('3. Entity Relationships', level=1)
    doc.add_paragraph("• [Pro Tier] --(INCLUDES)--> [Advanced Analytics]")
    doc.add_paragraph("• [Enterprise Tier] --(HAS_ACCESS_TO)--> [API]")
    doc.add_paragraph("• [Sarah Jenkins] --(MANAGES)--> [Enterprise Escalations]")
    doc.add_paragraph("• [Marcus Lee] --(RESOLVES)--> [API Bugs]")
    
    # Section 4: Standard Operating Procedures (SOPs)
    doc.add_heading('4. Standard Operating Procedures', level=1)
    
    doc.add_heading('4.1. Password Resets', level=2)
    doc.add_paragraph("To reset a password, the agent must ask the user for their company email. Once provided, instruct the user to check their email for a secure reset link. Do not ask for their current password.")
    
    doc.add_heading('4.2. Adding a New Employee', level=2)
    doc.add_paragraph("1. Navigate to the 'Team' tab on the ZenithHR dashboard.")
    doc.add_paragraph("2. Click 'Add Team Member'.")
    doc.add_paragraph("3. Enter the employee's First Name, Last Name, and Email.")
    doc.add_paragraph("4. Assign a role (Admin, Manager, or Employee).")
    
    doc.add_heading('4.3. Handling API Rate Limits', level=2)
    doc.add_paragraph("If a user encounters a '429 Too Many Requests' error, inform them that the standard API limit is 1,000 requests per minute. They must wait 60 seconds before retrying.")
    
    doc.save(filename)
    print(f"Successfully saved {filename}")

if __name__ == "__main__":
    create_docx("Startup_Agent_Knowledge.docx")
